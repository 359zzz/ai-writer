from __future__ import annotations

import re
import shutil
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Literal

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from markdown import markdown as md_to_html
from sqlmodel import select

from ..db import get_session
from ..models import Chapter, Project


router = APIRouter(prefix="/api/projects/{project_id}/export", tags=["export"])

ExportFormat = Literal["docx", "epub", "pdf"]


def _now_tag() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")


def _exports_dir() -> Path:
    api_root = Path(__file__).resolve().parents[2]  # .../apps/api
    out = api_root / "exports"
    out.mkdir(parents=True, exist_ok=True)
    return out


def _templates_dir() -> Path:
    api_root = Path(__file__).resolve().parents[2]  # .../apps/api
    out = api_root / "templates" / "export"
    out.mkdir(parents=True, exist_ok=True)
    return out


def _ensure_reference_docx() -> Path | None:
    """
    Create a pandoc reference.docx for nicer default styling (fonts/headings).

    This file is generated locally (not required for export to work).
    """
    try:
        ref = _templates_dir() / "reference.docx"
        if ref.exists():
            return ref
        from docx import Document
        from docx.shared import Pt

        doc = Document()

        # Base font (Windows-friendly; OK if missing on non-Windows).
        try:
            normal = doc.styles["Normal"]
            normal.font.name = "Microsoft YaHei"
            normal.font.size = Pt(11)
        except Exception:
            pass

        # Headings
        for style_name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)]:
            try:
                s = doc.styles[style_name]
                s.font.name = "Microsoft YaHei"
                s.font.size = Pt(size)
                s.font.bold = True
            except Exception:
                continue

        # Ensure at least one paragraph exists.
        doc.add_paragraph("")
        doc.save(ref)
        return ref
    except Exception:
        return None


def _epub_css_path() -> Path:
    return _templates_dir() / "epub.css"


def _compile_markdown(project: Project, chapters: list[Chapter]) -> str:
    # Pandoc-friendly markdown with metadata + explicit page breaks.
    # (DOCX/PDF can respect \\newpage; EPUB will simply ignore it.)
    now_local = datetime.now().strftime("%Y-%m-%d")
    safe_title = project.title.replace('"', "'")
    parts: list[str] = [
        "---",
        f'title: \"{safe_title}\"',
        f'date: \"{now_local}\"',
        "lang: zh-CN",
        "---",
        "",
        "\\newpage",
        "",
    ]
    for idx, ch in enumerate(chapters):
        if idx > 0:
            parts.append("")
            parts.append("\\newpage")
            parts.append("")
        parts.append(f"# {ch.chapter_index}. {ch.title}".strip())
        parts.append("")
        parts.append(ch.markdown.strip())
        parts.append("")
    return "\n".join(parts).strip() + "\n"


def _strip_markdown(md: str) -> str:
    # Very light markdown stripping (good enough for basic PDF fallback).
    t = re.sub(r"`{1,3}.*?`{1,3}", "", md, flags=re.S)
    t = re.sub(r"^#+\\s*", "", t, flags=re.M)
    t = re.sub(r"\\*\\*([^*]+)\\*\\*", r"\\1", t)
    t = re.sub(r"\\*([^*]+)\\*", r"\\1", t)
    t = re.sub(r"\\[(.*?)\\]\\((.*?)\\)", r"\\1 (\\2)", t)
    return t


def _pandoc_available() -> bool:
    return shutil.which("pandoc") is not None


def _pdf_engine() -> str | None:
    # Prefer XeLaTeX for CJK when available.
    for eng in ("xelatex", "lualatex", "pdflatex"):
        if shutil.which(eng):
            return eng
    return None


def _export_with_pandoc(md_path: Path, out_path: Path, fmt: ExportFormat) -> None:
    cmd: list[str] = [
        "pandoc",
        str(md_path),
        "-o",
        str(out_path),
        "--toc",
        "--toc-depth=2",
        "--number-sections",
        "--standalone",
    ]

    if fmt == "docx":
        ref = _ensure_reference_docx()
        if ref:
            cmd += ["--reference-doc", str(ref)]
    elif fmt == "epub":
        css = _epub_css_path()
        if css.exists():
            cmd += ["--css", str(css)]
        # Split by top-level headers (chapters).
        cmd += ["--split-level=1"]
    elif fmt == "pdf":
        eng = _pdf_engine()
        if not eng:
            raise RuntimeError("pdf_engine_missing")
        cmd += [
            "--pdf-engine",
            eng,
            "-V",
            "geometry:margin=1in",
            "-V",
            "fontsize=12pt",
            # Windows-friendly CJK default; ignored if font is missing.
            "-V",
            "mainfont=Microsoft YaHei",
        ]

    subprocess.run(cmd, check=True, capture_output=True)


def _export_docx_basic(md_text: str, out_path: Path) -> None:
    from docx import Document

    doc = Document()
    for line in md_text.splitlines():
        t = line.rstrip()
        if not t:
            doc.add_paragraph("")
            continue
        if t.startswith("# "):
            doc.add_heading(t[2:].strip(), level=1)
            continue
        if t.startswith("## "):
            doc.add_heading(t[3:].strip(), level=2)
            continue
        doc.add_paragraph(_strip_markdown(t))
    doc.save(out_path)


def _export_epub_basic(project: Project, md_text: str, out_path: Path) -> None:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier(project.id)
    book.set_title(project.title)
    book.set_language("zh")

    html = md_to_html(md_text, extensions=["fenced_code", "tables"])
    chapter = epub.EpubHtml(title=project.title, file_name="chap_001.xhtml", lang="zh")
    chapter.content = html
    book.add_item(chapter)

    book.toc = (chapter,)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    style = "body { font-family: serif; line-height: 1.6; }"
    nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
    book.add_item(nav_css)
    book.spine = ["nav", chapter]

    epub.write_epub(str(out_path), book, {})


def _export_pdf_basic(md_text: str, out_path: Path) -> None:
    # Basic text PDF fallback (no rich formatting).
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    # Try to register a common CJK font if present; otherwise fallback.
    # Do NOT fail export if font is missing.
    try:
        # Microsoft YaHei is common on Windows.
        font_path = Path("C:/Windows/Fonts/msyh.ttc")
        if font_path.exists():
            pdfmetrics.registerFont(TTFont("msyh", str(font_path)))
            font_name = "msyh"
        else:
            font_name = "Helvetica"
    except Exception:
        font_name = "Helvetica"

    text = _strip_markdown(md_text)
    c = canvas.Canvas(str(out_path), pagesize=LETTER)
    width, height = LETTER
    x = 40
    y = height - 50
    c.setFont(font_name, 11)
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if y < 60:
            c.showPage()
            c.setFont(font_name, 11)
            y = height - 50
        c.drawString(x, y, line[:120])
        y -= 14
    c.save()


@router.post("")
def export_project(project_id: str, payload: dict[str, Any]) -> FileResponse:
    fmt = str(payload.get("format") or "docx").lower()
    if fmt not in ("docx", "epub", "pdf"):
        raise HTTPException(status_code=400, detail="format must be docx|epub|pdf")

    with get_session() as session:
        project = session.get(Project, project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")

        chapters = list(
            session.exec(
                select(Chapter)
                .where(Chapter.project_id == project_id)
                .order_by(Chapter.chapter_index.asc(), Chapter.created_at.asc())
            )
        )
        if not chapters:
            raise HTTPException(status_code=400, detail="No chapters to export")

    md_text = _compile_markdown(project, chapters)
    out_dir = _exports_dir()
    base = f"{project.title}_{_now_tag()}".replace(" ", "_")
    md_path = out_dir / f"{base}.md"
    md_path.write_text(md_text, encoding="utf-8")

    out_path = out_dir / f"{base}.{fmt}"

    try:
        if _pandoc_available():
            try:
                _export_with_pandoc(md_path, out_path, fmt)  # type: ignore[arg-type]
            except RuntimeError as e:
                # PDF requires extra tooling (LaTeX engine). Fall back gracefully.
                if fmt == "pdf" and str(e) == "pdf_engine_missing":
                    _export_pdf_basic(md_text, out_path)
                else:
                    raise
        elif fmt == "docx":
            _export_docx_basic(md_text, out_path)
        elif fmt == "epub":
            _export_epub_basic(project, md_text, out_path)
        else:
            _export_pdf_basic(md_text, out_path)
    except subprocess.CalledProcessError as e:
        raise HTTPException(status_code=500, detail=f"pandoc_failed: {e.stderr[:200].decode('utf-8', 'ignore')}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"export_failed:{type(e).__name__}")

    return FileResponse(
        path=str(out_path),
        media_type="application/octet-stream",
        filename=out_path.name,
    )
